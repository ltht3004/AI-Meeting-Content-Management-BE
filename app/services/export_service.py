from io import BytesIO
from pathlib import Path
from xml.sax.saxutils import escape


def _value(value, fallback="N/A"):
    return value if value not in (None, "") else fallback


def _register_pdf_fonts() -> tuple[str, str]:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # Use Unicode fonts when available so Vietnamese names/locations render correctly in PDF.
    font_candidates = [
        (
            Path("C:/Windows/Fonts/arial.ttf"),
            Path("C:/Windows/Fonts/arialbd.ttf"),
            "ArialUnicode",
            "ArialUnicode-Bold"
        ),
        (
            Path("C:/Windows/Fonts/calibri.ttf"),
            Path("C:/Windows/Fonts/calibrib.ttf"),
            "CalibriUnicode",
            "CalibriUnicode-Bold"
        ),
        (
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
            "DejaVuSans",
            "DejaVuSans-Bold"
        ),
    ]

    for regular_path, bold_path, regular_name, bold_name in font_candidates:
        if regular_path.exists() and bold_path.exists():
            if regular_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(regular_name, str(regular_path)))
            if bold_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(bold_name, str(bold_path)))

            return regular_name, bold_name

    return "Helvetica", "Helvetica-Bold"


def _paragraph(text, style):
    from reportlab.platypus import Paragraph

    return Paragraph(escape(str(text)), style)


def generate_pdf_report(meeting_data: dict) -> BytesIO:
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    except ImportError as exc:
        raise RuntimeError("reportlab is required to export PDF reports.") from exc

    buffer = BytesIO()
    regular_font, bold_font = _register_pdf_fonts()
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "ReportTitle",
        parent=styles["Title"],
        fontName=bold_font,
        fontSize=20,
        leading=24,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#0f172a"),
        spaceAfter=8,
    )
    subtitle_style = ParagraphStyle(
        "ReportSubtitle",
        parent=styles["Normal"],
        fontName=regular_font,
        fontSize=10,
        leading=14,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#64748b"),
        spaceAfter=16,
    )
    section_style = ParagraphStyle(
        "SectionHeading",
        parent=styles["Heading2"],
        fontName=bold_font,
        fontSize=13,
        leading=16,
        textColor=colors.HexColor("#0052cc"),
        spaceBefore=10,
        spaceAfter=8,
    )
    body_style = ParagraphStyle(
        "BodyTextUnicode",
        parent=styles["BodyText"],
        fontName=regular_font,
        fontSize=10,
        leading=15,
        textColor=colors.HexColor("#1f2937"),
    )
    body_bold_style = ParagraphStyle(
        "BodyTextUnicodeBold",
        parent=body_style,
        fontName=bold_font,
    )
    bullet_style = ParagraphStyle(
        "BulletTextUnicode",
        parent=body_style,
        leftIndent=12,
        firstLineIndent=-8,
    )

    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=f"{_value(meeting_data.get('title'))} Report",
    )

    story = [
        _paragraph("AI Meeting Report", title_style),
        _paragraph(_value(meeting_data.get("title")), subtitle_style),
    ]

    info_rows = [
        [_paragraph("Date & Time", body_bold_style), _paragraph(_value(meeting_data.get("meeting_date")), body_style)],
        [_paragraph("Location", body_bold_style), _paragraph(_value(meeting_data.get("location")), body_style)],
        [_paragraph("Duration", body_bold_style), _paragraph(f"{_value(meeting_data.get('duration'))} minutes", body_style)],
        [_paragraph("Status", body_bold_style), _paragraph(_value(meeting_data.get("status")), body_style)],
    ]
    info_table = Table(info_rows, colWidths=[34 * mm, 126 * mm])
    info_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#eff6ff")),
        ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#0f172a")),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#dbeafe")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))

    story.extend([
        _paragraph("Meeting Information", section_style),
        info_table,
        _paragraph("Description", section_style),
        _paragraph(_value(meeting_data.get("description"), "No description."), body_style),
        _paragraph("Participants", section_style),
    ])

    participants = meeting_data.get("participants") or []
    if participants:
        if len(participants) > 8:
            # Long participant lists are split into 3 columns to avoid wasting vertical space.
            row_count = (len(participants) + 2) // 3
            columns = [
                participants[0:row_count],
                participants[row_count:row_count * 2],
                participants[row_count * 2:],
            ]
            participant_rows = []

            for index in range(row_count):
                participant_rows.append([
                    _paragraph(f"- {column[index]}" if index < len(column) else "", bullet_style)
                    for column in columns
                ])

            participant_table = Table(participant_rows, colWidths=[53 * mm, 53 * mm, 53 * mm])
            participant_table.setStyle(TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ]))
            story.append(participant_table)
        else:
            story.extend([_paragraph(f"- {participant}", bullet_style) for participant in participants])
    else:
        story.append(_paragraph("No participants.", body_style))

    story.extend([
        _paragraph("AI Summary", section_style),
        _paragraph(_value(meeting_data.get("summary"), "No summary generated yet."), body_style),
        _paragraph("Recordings", section_style),
    ])

    recordings = meeting_data.get("recordings") or []
    if recordings:
        for recording in recordings:
            story.append(_paragraph(f"- {_value(recording.get('file_name'))} ({_value(recording.get('size_label'))})", bullet_style))
    else:
        story.append(_paragraph("No recordings uploaded.", body_style))

    story.append(_paragraph("Transcript", section_style))
    transcripts = meeting_data.get("transcripts") or []
    if transcripts:
        for item in transcripts:
            story.append(_paragraph(_value(item.get("recording_name")), body_bold_style))
            story.append(_paragraph(_value(item.get("content"), "No transcript content."), body_style))
            story.append(Spacer(1, 6))
    else:
        story.append(_paragraph("No transcript available.", body_style))

    document.build(story)
    buffer.seek(0)
    return buffer


def generate_docx_report(meeting_data: dict) -> BytesIO:
    try:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError("python-docx is required to export Word reports.") from exc

    def set_cell_shading(cell, fill: str):
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        tc_pr.append(shading)

    def set_cell_border(cell, color: str = "DBEAFE"):
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.first_child_found_in("w:tcBorders")
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)

        for edge in ("top", "left", "bottom", "right"):
            tag = f"w:{edge}"
            element = borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                borders.append(element)
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "4")
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), color)

    def format_run(run, size=10, bold=False, color="1F2937"):
        # Force Arial for Latin and East Asian text so Word displays Vietnamese consistently.
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor.from_string(color)

    def add_section(title: str):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run(title)
        format_run(run, size=12, bold=True, color="0052CC")
        return paragraph

    def add_body(text: str):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.15
        format_run(paragraph.add_run(str(text)), size=10)
        return paragraph

    def add_bullet(text: str):
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        format_run(paragraph.add_run(str(text)), size=10)
        return paragraph

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal_style.font.size = Pt(10)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("AI Meeting Report")
    format_run(title_run, size=20, bold=True, color="0F172A")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run(_value(meeting_data.get("title")))
    format_run(subtitle_run, size=10, color="64748B")

    divider = document.add_paragraph()
    divider.paragraph_format.space_after = Pt(10)
    divider_run = divider.add_run("")
    divider_run._r.get_or_add_rPr()
    p_pr = divider._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0284C7")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    add_section("Meeting Information")
    info_table = document.add_table(rows=0, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = False
    for label, value in [
        ("Date & Time", meeting_data.get("meeting_date")),
        ("Location", meeting_data.get("location")),
        ("Duration", f"{_value(meeting_data.get('duration'))} minutes"),
        ("Status", meeting_data.get("status")),
    ]:
        row = info_table.add_row().cells
        row[0].width = Inches(1.45)
        row[1].width = Inches(5.4)
        row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        row[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_shading(row[0], "EFF6FF")
        set_cell_border(row[0])
        set_cell_border(row[1])
        row[0].paragraphs[0].paragraph_format.space_after = Pt(0)
        row[1].paragraphs[0].paragraph_format.space_after = Pt(0)
        format_run(row[0].paragraphs[0].add_run(label), bold=True, color="0F172A")
        format_run(row[1].paragraphs[0].add_run(str(_value(value))))

    add_section("Description")
    add_body(_value(meeting_data.get("description"), "No description."))

    add_section("Participants")
    participants = meeting_data.get("participants") or []
    if participants:
        if len(participants) > 8:
            # Match the PDF layout: long participant lists are rendered in 3 columns.
            row_count = (len(participants) + 2) // 3
            columns = [
                participants[0:row_count],
                participants[row_count:row_count * 2],
                participants[row_count * 2:],
            ]
            participant_table = document.add_table(rows=row_count, cols=3)
            participant_table.autofit = False

            for index, row in enumerate(participant_table.rows):
                for cell_index, column in enumerate(columns):
                    participant = column[index] if index < len(column) else ""
                    cell = row.cells[cell_index]
                    cell.width = Inches(2.2)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                    paragraph = cell.paragraphs[0]
                    paragraph.paragraph_format.space_after = Pt(2)
                    if participant:
                        format_run(paragraph.add_run(f"- {participant}"), size=10)
        else:
            for participant in participants:
                add_bullet(participant)
    else:
        add_body("No participants.")

    add_section("AI Summary")
    add_body(_value(meeting_data.get("summary"), "No summary generated yet."))

    add_section("Recordings")
    recordings = meeting_data.get("recordings") or []
    if recordings:
        for recording in recordings:
            add_bullet(
                f"{_value(recording.get('file_name'))} ({_value(recording.get('size_label'))})",
            )
    else:
        add_body("No recordings uploaded.")

    add_section("Transcript")
    transcripts = meeting_data.get("transcripts") or []
    if transcripts:
        for item in transcripts:
            recording_name = document.add_paragraph()
            recording_name.paragraph_format.space_before = Pt(4)
            recording_name.paragraph_format.space_after = Pt(2)
            format_run(recording_name.add_run(_value(item.get("recording_name"))), bold=True, color="0284C7")
            add_body(_value(item.get("content"), "No transcript content."))
    else:
        add_body("No transcript available.")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
